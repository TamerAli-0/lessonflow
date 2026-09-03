from __future__ import annotations

import re
import shutil
import zipfile
from copy import deepcopy
from datetime import date, timedelta
from pathlib import Path
from typing import Iterable

import fitz
from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .schemas import LESSON_LIST_FIELDS, PlanningSettings


class TemplateError(RuntimeError):
    """Raised when required semantic anchors cannot be found in a template."""


STAGES = {
    "warm_up": ("warm up", "warm_up"),
    "focused_instruction": ("focused instruction", "focused"),
    "guided_instruction": ("guided instruction", "guided"),
    "collaborative_learning": ("collaborative learning", "collaborative"),
    "independent_learning": ("independent learning", "independent"),
    "progress_check": ("progress check", "progress_check"),
}

FIELD_CAPACITY = {
    "title": 140,
    "resources": 260,
    "kpis": 420,
    "c21_skills": 360,
    "success_criteria": 420,
    "cross_curricular": 280,
    "warm_up": 220,
    "focused_instruction": 420,
    "guided_instruction": 420,
    "collaborative_learning": 380,
    "independent_learning": 340,
    "progress_check": 340,
    "home_learning": 180,
    "reflection": 320,
}


def export_lesson_bundle(
    template_path: str | Path,
    output_path: str | Path,
    analysis: dict,
    plan: dict,
    image_assets: list[dict] | None = None,
    force_opening_page_break: bool = True,
) -> Path:
    template = Path(template_path)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    settings = PlanningSettings.from_dict(plan["settings"])
    lesson_files: list[Path] = []
    lessons_dir = output.parent / "lesson-documents"
    if lessons_dir.exists():
        shutil.rmtree(lessons_dir)
    lessons_dir.mkdir(parents=True)

    for lesson in plan["lessons"]:
        filename = f"{lesson['lesson_number']:02d}-{safe_filename(lesson['title'])}.docx"
        destination = lessons_dir / filename
        shutil.copyfile(template, destination)
        fill_lesson_document(
            destination, analysis, lesson, settings, plan.get("options", {}), image_assets or [],
            force_opening_page_break=force_opening_page_break,
        )
        lesson_files.append(destination)

    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for file_path in lesson_files:
            archive.write(file_path, arcname=file_path.name)
    return output


def fill_lesson_document(
    path: str | Path,
    analysis: dict,
    lesson: dict,
    settings: PlanningSettings,
    options: dict | None = None,
    image_assets: list[dict] | None = None,
    force_opening_page_break: bool = True,
) -> None:
    document = Document(path)
    if not document.tables:
        raise TemplateError("The template does not contain a lesson-plan table.")
    main_table = find_main_table(document.tables)
    options = options or {}
    prepared_images = prepare_lesson_images(
        document,
        main_table,
        lesson,
        image_assets or [],
        str(lesson.get("recommended_image_stage") or options.get("image_placement", "focused_instruction")),
        int(options.get("max_images_per_lesson", 0)) if options.get("extract_images") else 0,
    )

    set_neighbor_value(main_table, "course", analysis.get("course_title") or analysis.get("module_title"))
    set_neighbor_value(main_table, "class", settings.student_level)
    set_week_fields(main_table, lesson["lesson_number"], settings)
    set_topic_overview(main_table, analysis, lesson)

    replace_section(main_table, "main learning focus", lesson.get("kpis", []), FIELD_CAPACITY["kpis"])
    replace_section(main_table, "c21st skills", lesson.get("c21_skills", []), FIELD_CAPACITY["c21_skills"])
    replace_section(main_table, "success criteria", lesson.get("success_criteria", []), FIELD_CAPACITY["success_criteria"])
    replace_section(main_table, "numeracy, literacy", lesson.get("cross_curricular", []), FIELD_CAPACITY["cross_curricular"])

    for field, (stage_anchor, timing_key) in STAGES.items():
        content_cell, time_cell, extra_cells = find_stage_slot(main_table, stage_anchor)
        replace_cell_content(content_cell, lesson.get(field, []), bullets=True, capacity=FIELD_CAPACITY[field])
        timing = lesson.get("timings", {}).get(timing_key)
        if time_cell is not None and timing is not None:
            replace_cell_content(time_cell, [str(timing)], bullets=False)
        if field == "progress_check":
            homework_cell = next((cell for cell in extra_cells if "home learning" in normalize(cell.text)), None)
            if homework_cell is not None:
                replace_after_heading(
                    homework_cell,
                    lesson.get("home_learning", []),
                    use_existing_body=False,
                    capacity=FIELD_CAPACITY["home_learning"],
                )

    compact_activity_page_for_density(main_table, lesson)
    restore_lesson_images(prepared_images)

    replace_reflection(document.tables, lesson.get("reflection", ""), FIELD_CAPACITY["reflection"])
    apply_lesson_text_formatting(document, lesson)
    stabilize_pagination(main_table, document.tables, allow_flow=not force_opening_page_break)
    if force_opening_page_break:
        # The designed two-page layout starts page 2 at Opening/Warm Up. When the plan has grown
        # past two pages that hard break only strands a near-empty page, so the caller re-runs
        # without it and the content flows continuously instead.
        split_table_before_opening(main_table)
    compact_trailing_paragraphs(document)
    document.save(path)


def find_main_table(tables: Iterable[Table]) -> Table:
    for table in tables:
        text = normalize(" ".join(cell.text for row in table.rows for cell in unique_cells(row)))
        if "warm up" in text and "guided instruction" in text:
            return table
    raise TemplateError("Could not find the table containing the fixed lesson stages.")


def unique_cells(row) -> list[_Cell]:
    seen = set()
    cells: list[_Cell] = []
    for cell in row.cells:
        if cell._tc in seen:
            continue
        seen.add(cell._tc)
        cells.append(cell)
    return cells


def normalize(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(text).lower()).strip()


def set_neighbor_value(table: Table, label: str, value: str) -> None:
    anchor = normalize(label)
    for row in table.rows:
        cells = unique_cells(row)
        for index, cell in enumerate(cells[:-1]):
            if normalize(cell.text) == anchor:
                replace_cell_content(cells[index + 1], [str(value or "")], bullets=False)
                return


def set_week_fields(table: Table, lesson_number: int, settings: PlanningSettings) -> None:
    start = date.fromisoformat(settings.starting_date)
    end = start + timedelta(days=(settings.module_weeks * 7) - 1)
    week_label = "W1" if settings.module_weeks == 1 else f"W1-W{settings.module_weeks}"
    if start.year == end.year and start.month == end.month:
        date_label = f"({start.day}-{end.day}) {start.strftime('%b %Y')}"
    elif start.year == end.year:
        date_label = f"{start.strftime('%d %b')}-{end.strftime('%d %b %Y')}"
    else:
        date_label = f"{start.strftime('%d %b %y')}-{end.strftime('%d %b %y')}"
    for row in table.rows:
        cells = unique_cells(row)
        for index, cell in enumerate(cells):
            if "week date" in normalize(cell.text):
                values = cells[index + 1 :]
                if values:
                    replace_cell_content(values[0], [week_label], bullets=False)
                if len(values) > 1:
                    replace_cell_content(values[1], [date_label], bullets=False)
                return


def set_topic_overview(table: Table, analysis: dict, lesson: dict) -> None:
    targets = []
    for row in table.rows:
        for cell in unique_cells(row):
            if "lesson unit topic" in normalize(cell.text) and "resources" in normalize(cell.text):
                targets.append(cell)
    if not targets:
        raise TemplateError("Could not find the lesson/topic/resources area in the template.")

    seen = set()
    for target in targets:
        if target._tc in seen:
            continue
        seen.add(target._tc)
        nonempty = [paragraph for paragraph in target.paragraphs if paragraph.text.strip()]
        if nonempty:
            set_paragraph_text(nonempty[0], analysis.get("module_title", "Module"))
            adapt_paragraph_font(nonempty[0], len(str(analysis.get("module_title", ""))), 110)
        source_reference = compact_source_reference(str(lesson.get("page_range", "")))
        topic_value = f"{lesson['title']} ({source_reference})" if source_reference else str(lesson["title"])
        topic_paragraph = replace_value_after_anchor(target, "lesson unit topic", topic_value)
        resource_value = "; ".join(lesson.get("resources", []))
        resource_paragraph = replace_value_after_anchor(target, "resources", resource_value)
        if topic_paragraph is not None:
            adapt_paragraph_font(topic_paragraph, len(topic_value), FIELD_CAPACITY["title"])
        if resource_paragraph is not None:
            adapt_paragraph_font(resource_paragraph, len(resource_value), FIELD_CAPACITY["resources"])


def compact_source_reference(value: str) -> str:
    reference = " ".join(str(value).split())
    match = re.fullmatch(r"document\s+parts?\s+(\d+)\s*[-–]\s*\1", reference, flags=re.I)
    if match:
        return f"document part {match.group(1)}"
    return reference


def replace_value_after_anchor(cell: _Cell, anchor: str, value: str) -> Paragraph | None:
    normalized_anchor = normalize(anchor)
    for index, paragraph in enumerate(cell.paragraphs):
        if normalized_anchor in normalize(paragraph.text):
            for following in cell.paragraphs[index + 1 :]:
                if following.text.strip():
                    set_paragraph_text(following, value)
                    return following
    return None


def replace_section(table: Table, anchor: str, lines: list[str], capacity: int) -> None:
    normalized_anchor = normalize(anchor)
    for row in table.rows:
        for cell in unique_cells(row):
            text = normalize(cell.text)
            if normalized_anchor in text:
                replace_after_heading(cell, lines, capacity=capacity)
                return
    raise TemplateError(f"Could not locate the '{anchor}' section in the template.")


def replace_after_heading(
    cell: _Cell,
    lines: list[str],
    *,
    use_existing_body: bool = True,
    capacity: int | None = None,
) -> None:
    nonempty = [paragraph for paragraph in cell.paragraphs if paragraph.text.strip()]
    heading = nonempty[0].text if nonempty else ""
    heading_p_pr, heading_r_pr = paragraph_format(nonempty[0] if nonempty else cell.paragraphs[0])
    if use_existing_body and len(nonempty) > 1:
        body_p_pr, body_r_pr = paragraph_format(nonempty[1])
    else:
        body_p_pr, body_r_pr = plain_body_format(cell)
    # The heading keeps the template's emphasis because it is a fixed label. Generated body text
    # must not, or the editor's bold toggle would disagree with what the teacher sees.
    body_r_pr = _without_emphasis(body_r_pr)
    rebuild_cell(
        cell,
        [(heading, heading_p_pr, heading_r_pr, False)]
        + [(str(line), body_p_pr, body_r_pr, use_existing_body) for line in lines if str(line).strip()],
    )
    if capacity:
        adapt_cell_font(cell, sum(len(str(line)) for line in lines), capacity, skip_paragraphs=1)


def find_stage_slot(table: Table, stage_anchor: str) -> tuple[_Cell, _Cell | None, list[_Cell]]:
    matches: list[tuple[int, _Cell, _Cell | None, list[_Cell]]] = []
    anchor = normalize(stage_anchor)
    for row_index, row in enumerate(table.rows):
        cells = unique_cells(row)
        for index, cell in enumerate(cells):
            if anchor in normalize(cell.text):
                right = cells[index + 1 :]
                if not right:
                    continue
                time_cell = cells[index - 1] if index > 0 else None
                content = right[0]
                matches.append((len(content.text), content, time_cell, right[1:]))
    if not matches:
        raise TemplateError(f"Could not locate the fixed stage '{stage_anchor}'.")
    _, content, time_cell, extra = max(matches, key=lambda item: item[0])
    return content, time_cell, extra


def replace_reflection(tables: Iterable[Table], reflection: str, capacity: int) -> None:
    for table in tables:
        all_text = normalize(" ".join(cell.text for row in table.rows for cell in unique_cells(row)))
        if "reflection" not in all_text or "recap success criteria" not in all_text:
            continue
        last_cells = unique_cells(table.rows[-1])
        if last_cells:
            replace_cell_content(last_cells[-1], [reflection], bullets=False, capacity=capacity)
            return
    raise TemplateError("Could not locate the reflection area in the template.")


def replace_cell_content(
    cell: _Cell,
    lines: Iterable[str],
    *,
    bullets: bool,
    heading_count: int = 0,
    capacity: int | None = None,
) -> None:
    clean_lines = [str(line).strip() for line in lines if str(line).strip()]
    if not clean_lines:
        clean_lines = [""]
    base_p_pr, base_r_pr = plain_body_format(cell)
    rebuild_cell(
        cell,
        [
            (line, base_p_pr, base_r_pr, bullets and index >= heading_count)
            for index, line in enumerate(clean_lines)
        ],
    )
    if capacity:
        adapt_cell_font(cell, sum(len(line) for line in clean_lines), capacity, skip_paragraphs=heading_count)


def rebuild_cell(cell: _Cell, rows: list[tuple[str, object, object, bool]]) -> None:
    for paragraph in list(cell.paragraphs[1:]):
        paragraph._element.getparent().remove(paragraph._element)
    first = cell.paragraphs[0]
    first.clear()
    for index, (text, p_pr, r_pr, bullet) in enumerate(rows or [("", None, None, False)]):
        paragraph = first if index == 0 else cell.add_paragraph()
        if p_pr is not None:
            if paragraph._p.pPr is not None:
                paragraph._p.remove(paragraph._p.pPr)
            paragraph._p.insert(0, deepcopy(p_pr))
        if bullet:
            if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
                paragraph._p.pPr.remove(paragraph._p.pPr.numPr)
            try:
                paragraph.style = "List Bullet"
            except KeyError:
                pass
        run = paragraph.add_run(text)
        if r_pr is not None:
            run._r.insert(0, deepcopy(r_pr))
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0


def paragraph_format(paragraph: Paragraph) -> tuple[object | None, object | None]:
    p_pr = deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None
    run = next((item for item in paragraph.runs if item.text.strip()), None)
    r_pr = deepcopy(run._r.rPr) if run is not None and run._r.rPr is not None else None
    return p_pr, r_pr


def plain_body_format(cell: _Cell) -> tuple[object | None, object | None]:
    candidates: list[tuple[int, Paragraph, object]] = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            text = run.text.strip()
            if len(re.findall(r"[A-Za-zÀ-ÿ]", text)) < 4:
                continue
            score = min(len(text), 40)
            score += 30 if run.font.color.rgb is None else 0
            score += 20 if not run.italic else 0
            score += 10 if not run.bold else 0
            candidates.append((score, paragraph, run))
    if candidates:
        _, paragraph, run = max(candidates, key=lambda item: item[0])
        p_pr = deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None
        r_pr = deepcopy(run._r.rPr) if run._r.rPr is not None else None
        return p_pr, _without_emphasis(r_pr)

    paragraph = cell.paragraphs[0]
    p_pr = deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    r_pr = deepcopy(run._r.rPr) if run._r.rPr is not None else None
    return p_pr, _without_emphasis(r_pr)


def _without_emphasis(r_pr):
    """Strip inherited bold/italic from the template run generated text is modelled on.

    Whether the sampled template run happened to be bold decided whether whole cells came out
    bold, which the editor could not know about: its toggle reads lesson.text_formatting only.
    Clearing the two properties the editor controls keeps what is shown and what is stored in
    agreement. Font, size, and colour are still inherited so the template still looks itself.
    """
    if r_pr is None:
        return None
    for tag in ("b", "bCs", "i", "iCs"):
        for node in r_pr.findall(qn(f"w:{tag}")):
            r_pr.remove(node)
    return r_pr


def prepare_lesson_images(
    document: Document,
    table: Table,
    lesson: dict,
    image_assets: list[dict],
    placement: str,
    maximum: int,
) -> list[dict]:
    if maximum <= 0:
        return []
    allowed_fields = {field: stage_anchor for field, (stage_anchor, _) in STAGES.items()}
    assets_by_xref = {int(asset.get("xref", -1)): asset for asset in image_assets}
    default_field = placement if placement in allowed_fields else "focused_instruction"

    if "selected_images" in lesson:
        selections = lesson.get("selected_images", [])
    elif "selected_image_xrefs" in lesson:
        selections = [{"xref": xref, "field": default_field} for xref in lesson.get("selected_image_xrefs", [])]
    else:
        recommended_pages = {int(page) for page in lesson.get("recommended_image_pages", []) if str(page).isdigit()}
        candidates = [asset for asset in image_assets if int(asset.get("page", 0)) in recommended_pages]
        candidates.sort(key=lambda asset: int(asset.get("area", 0)), reverse=True)
        selections = [{"xref": candidates[0]["xref"], "field": default_field}] if candidates else []

    prepared: list[dict] = []
    used_xrefs: set[int] = set()
    used_fields: set[str] = set()
    for selection in selections:
        try:
            xref = int(selection.get("xref"))
        except (AttributeError, TypeError, ValueError):
            continue
        field = str(selection.get("field") or default_field)
        if field not in allowed_fields or xref in used_xrefs or field in used_fields:
            continue
        asset = assets_by_xref.get(xref)
        if asset is None or not Path(str(asset.get("path", ""))).exists():
            continue
        section_value = lesson.get(field, [])
        section_items = section_value if isinstance(section_value, list) else [str(section_value)]
        section_text_length = sum(len(str(value)) for value in section_items)
        # A density guard used to withhold an image from a text-heavy section. It made the editor
        # claim an image was selected while the page showed none, and it only lifted once the image
        # happened to gain a position. The teacher decides: a selected image is always placed, and
        # the two-page render check before download remains the real guard against overflow.
        content_cell, _, _ = find_stage_slot(table, allowed_fields[field])
        anchors = capture_cell_image_anchors(document, content_cell)
        prepared.append(
            {
                "cell": content_cell,
                "asset": asset,
                "anchors": anchors,
                "field": field,
                "placement": selection,
                "section_text_length": section_text_length,
                "section_item_count": len([value for value in section_items if str(value).strip()]),
            }
        )
        used_xrefs.add(xref)
        used_fields.add(field)
        if len(prepared) >= maximum:
            break
    for item in prepared:
        item["lesson_image_count"] = len(prepared)
    return prepared


def capture_cell_image_anchors(document: Document, content_cell: _Cell) -> list[tuple[object, object]]:
    anchors = []
    for drawing in content_cell._tc.xpath(".//w:drawing"):
        anchor_elements = drawing.xpath(".//wp:anchor")
        blips = drawing.xpath(".//a:blip")
        if not anchor_elements or not blips:
            continue
        relation_id = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        image_part = document.part.related_parts.get(relation_id)
        if image_part is None or not str(getattr(image_part, "content_type", "")).startswith("image/"):
            continue
        anchors.append((deepcopy(drawing), image_part))
    return anchors


def restore_lesson_images(prepared: list[dict]) -> None:
    for item in prepared:
        content_cell = item["cell"]
        asset = item["asset"]
        anchors = item.get("anchors", [])
        compact_text_for_image_position(content_cell, item.get("placement", {}))
        if anchors:
            paragraph = content_cell.paragraphs[-1]
            run = paragraph.add_run()
            drawing, image_part = anchors[0]
            image_part._blob = convert_image_for_part(Path(asset["path"]), image_part.content_type)
            resize_floating_image(
                drawing,
                int(asset.get("width", 1)),
                int(asset.get("height", 1)),
                item.get("placement", {}),
                int(item.get("section_text_length", 0)),
                int(item.get("lesson_image_count", 1)),
                int(item.get("section_item_count", 1)),
            )
            run._r.append(drawing)
        else:
            paragraph = content_cell.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            text_length = int(item.get("section_text_length", 0))
            inline_width = 1.70 if text_length <= 140 else 1.45 if text_length <= 260 else 1.20
            run.add_picture(str(asset["path"]), width=Inches(inline_width))


def convert_image_for_part(path: Path, content_type: str) -> bytes:
    pixmap = fitz.Pixmap(str(path))
    if pixmap.colorspace is None or pixmap.colorspace.n not in (1, 3):
        pixmap = fitz.Pixmap(fitz.csRGB, pixmap)
    if content_type in {"image/jpeg", "image/jpg"}:
        return pixmap.tobytes("jpeg")
    return pixmap.tobytes("png")


def compact_text_for_image_position(cell: _Cell, placement: dict | None) -> None:
    placement = placement or {}
    try:
        x_ratio = min(1.0, max(0.0, float(placement.get("x_ratio", 1.0))))
    except (TypeError, ValueError):
        x_ratio = 1.0
    if x_ratio >= 0.78:
        return
    target_size = 8.5
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            current = run.font.size.pt if run.font.size is not None else 11.0
            if current > target_size:
                run.font.size = Pt(target_size)


def resize_floating_image(
    drawing,
    width: int,
    height: int,
    placement: dict | None = None,
    section_text_length: int = 0,
    lesson_image_count: int = 1,
    section_item_count: int = 1,
) -> None:
    anchor_nodes = drawing.xpath(".//wp:anchor")
    if not anchor_nodes:
        return
    anchor = anchor_nodes[0]
    anchor.set("behindDoc", "0")
    anchor.set("allowOverlap", "0")
    anchor.set("distL", str(6 * 12_700))
    anchor.set("distR", str(6 * 12_700))
    anchor.set("distT", str(2 * 12_700))
    anchor.set("distB", str(2 * 12_700))
    for child in list(anchor):
        if child.tag.rsplit("}", 1)[-1] in {"sizeRelH", "sizeRelV"}:
            anchor.remove(child)

    if section_text_length <= 140:
        max_width_inches, max_height_inches = 2.15, 1.25
    elif section_text_length <= 260:
        max_width_inches, max_height_inches = 1.80, 1.00
    else:
        max_width_inches, max_height_inches = 1.45, 0.78
    if section_item_count >= 3:
        max_width_inches = min(max_width_inches, 1.35)
        max_height_inches = min(max_height_inches, 0.78)
    elif section_item_count == 2:
        max_width_inches = min(max_width_inches, 1.65)
        max_height_inches = min(max_height_inches, 0.92)
    if lesson_image_count > 1:
        max_width_inches = min(max_width_inches, 1.50)
        max_height_inches = min(max_height_inches, 0.72)
    max_width = int(max_width_inches * 914_400)
    max_height = int(max_height_inches * 914_400)
    scale = min(max_width / max(width, 1), max_height / max(height, 1))
    cx = max(1, int(width * scale))
    cy = max(1, int(height * scale))
    for extent in drawing.xpath(".//wp:extent") + drawing.xpath(".//a:xfrm/a:ext"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))
    placement = placement or {}
    try:
        x_ratio = min(1.0, max(0.0, float(placement.get("x_ratio", 1.0))))
        y_ratio = min(1.0, max(0.0, float(placement.get("y_ratio", 0.58))))
    except (TypeError, ValueError):
        x_ratio, y_ratio = 1.0, 0.58
    vertical_offsets = drawing.xpath(".//wp:positionV/wp:posOffset")
    if vertical_offsets:
        if "offset_y_points" in placement:
            try:
                compaction_correction = 2.0 if x_ratio < 0.78 else 0.0
                vertical_points = float(placement["offset_y_points"]) + compaction_correction
                vertical_offsets[0].text = str(int(vertical_points * 12_700))
            except (TypeError, ValueError):
                vertical_offsets[0].text = "0"
        else:
            vertical_inches = -0.35 + (y_ratio * 0.60)
            vertical_offsets[0].text = str(int(vertical_inches * 914_400))
    horizontal_offsets = drawing.xpath(".//wp:positionH/wp:posOffset")
    if horizontal_offsets:
        base_horizontal = 6.15 * 914_400
        if "offset_x_points" in placement:
            try:
                horizontal_offsets[0].text = str(int(base_horizontal + (float(placement["offset_x_points"]) * 12_700)))
            except (TypeError, ValueError):
                horizontal_offsets[0].text = str(int(base_horizontal))
        else:
            horizontal_inches = 0.40 + (x_ratio * 5.75)
            horizontal_offsets[0].text = str(int(horizontal_inches * 914_400))
    wrap_nodes = drawing.xpath(".//wp:wrapTight") + drawing.xpath(".//wp:wrapSquare")
    wrap_side = "left" if x_ratio >= 0.5 else "right"
    for wrap in wrap_nodes:
        wrap.set("wrapText", wrap_side)


def adapt_cell_font(cell: _Cell, character_count: int, capacity: int, *, skip_paragraphs: int = 0) -> None:
    effective_capacity = max(capacity, estimate_cell_capacity(cell))
    size = adaptive_font_size(character_count, effective_capacity)
    if size is None:
        return
    for paragraph in cell.paragraphs[skip_paragraphs:]:
        for run in paragraph.runs:
            if run.text.strip():
                run.font.size = Pt(size)


def adapt_paragraph_font(paragraph: Paragraph, character_count: int, capacity: int) -> None:
    size = adaptive_font_size(character_count, capacity)
    if size is None:
        return
    for run in paragraph.runs:
        if run.text.strip():
            run.font.size = Pt(size)


def adaptive_font_size(character_count: int, capacity: int) -> float | None:
    ratio = character_count / max(capacity, 1)
    if ratio <= 1.0:
        return None
    if ratio <= 1.2:
        return 10.0
    if ratio <= 1.5:
        return 9.25
    return 8.5


def compact_activity_page_for_density(table: Table, lesson: dict) -> None:
    activity_fields = list(STAGES)
    density = sum(
        len(str(item))
        for field in activity_fields
        for item in (lesson.get(field, []) if isinstance(lesson.get(field, []), list) else [lesson.get(field, "")])
    )
    density += sum(len(str(item)) for item in lesson.get("home_learning", []))
    density += len(str(lesson.get("reflection", "")))
    if density <= 900:
        return
    target_size = 9.25 if density <= 1_450 else 8.5
    for field, (stage_anchor, _) in STAGES.items():
        content_cell, _, extra_cells = find_stage_slot(table, stage_anchor)
        cells = [content_cell]
        if field == "progress_check":
            cells.extend(cell for cell in extra_cells if "home learning" in normalize(cell.text))
        for cell in cells:
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if paragraph_index == 0:
                    continue
                for run in paragraph.runs:
                    if run.text.strip():
                        current = run.font.size.pt if run.font.size is not None else 11.0
                        if current > target_size:
                            run.font.size = Pt(target_size)


def apply_lesson_text_formatting(document: Document, lesson: dict) -> None:
    formatting = lesson.get("text_formatting", {})
    if not isinstance(formatting, dict):
        return
    paragraphs = [
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in unique_cells(row)
        for paragraph in cell.paragraphs
    ]
    used: set[int] = set()
    for field, specs in formatting.items():
        if not isinstance(specs, list):
            continue
        raw_value = lesson.get(field, "")
        values = raw_value if isinstance(raw_value, list) else [raw_value]
        for index, spec in enumerate(specs):
            if index >= len(values) or not isinstance(spec, dict):
                continue
            target = normalize(str(values[index]))
            if not target:
                continue
            paragraph = next(
                (item for item in paragraphs if id(item._p) not in used and normalize(item.text) == target),
                None,
            )
            if paragraph is None:
                continue
            used.add(id(paragraph._p))
            bold = spec.get("bold") if isinstance(spec.get("bold"), bool) else None
            italic = spec.get("italic") if isinstance(spec.get("italic"), bool) else None
            try:
                font_size = float(spec.get("font_size")) if spec.get("font_size") else None
            except (TypeError, ValueError):
                font_size = None
            if font_size is not None and not 8.0 <= font_size <= 18.0:
                font_size = None
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if font_size is not None:
                    run.font.size = Pt(font_size)


def estimate_cell_capacity(cell: _Cell) -> int:
    width = cell.width.pt if cell.width is not None else 0
    row = cell._tc.getparent()
    heights = row.xpath("./w:trPr/w:trHeight/@w:val") if row is not None else []
    if not width or not heights:
        return 0
    try:
        minimum_height = max(float(value) / 20.0 for value in heights)
    except (TypeError, ValueError):
        return 0
    if minimum_height < 35:
        return 0
    usable_width = max(width - 12, 20)
    usable_height = max(minimum_height - 18, 12)
    return int((usable_width / 4.2) * (usable_height / 11.0))


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    base_r_pr = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(str(text))
    if base_r_pr is not None:
        run._r.insert(0, base_r_pr)


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-.")
    return (cleaned or "lesson")[:80]


def stabilize_pagination(main_table: Table, tables: Iterable[Table], allow_flow: bool = False) -> None:
    """`allow_flow` is used once the plan has grown past the designed two pages.

    The rules that keep a two-page plan tidy - rows that never split, headers pinned to their
    content - are what push a short block such as Reflection onto a page of its own. When the
    document is already longer, letting everything flow fills each page to the bottom instead.
    """
    stage_anchors = (
        "warm up",
        "focused instruction",
        "guided instruction",
        "collaborative learning",
        "independent learning",
        "progress check",
    )
    rows = list(main_table.rows)
    for row_position, row in enumerate(rows):
        row_text = normalize(" ".join(cell.text for cell in unique_cells(row)))
        following = rows[row_position + 1] if row_position + 1 < len(rows) else None
        following_text = (
            normalize(" ".join(cell.text for cell in unique_cells(following))) if following is not None else ""
        )
        # Pinning a header to content that is itself too tall for the space pushes both to the
        # next page and strands a near-empty one. Let the header stay put and the content flow.
        pin_to_next = not allow_flow and len(following_text) <= ROW_SPLIT_THRESHOLD
        # A row that cannot split jumps whole to the next page when it does not fit, which leaves
        # a half-empty page behind. Small rows still stay intact; a tall one is allowed to flow so
        # an expanded plan fills its pages instead of stranding gaps.
        if allow_flow or len(row_text) > ROW_SPLIT_THRESHOLD:
            allow_row_split(row)
        else:
            set_row_cant_split(row)
        if any(anchor in row_text for anchor in stage_anchors):
            clear_row_minimum_height(row)
        if "how will you engage students in learning" in row_text:
            for cell in unique_cells(row):
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = pin_to_next
        if any(
            prompt in row_text
            for prompt in (
                "teacher directed activities",
                "student directed activities need help",
                "student directed activities want a challenge",
                "independent practice align questions",
            )
        ):
            for cell in unique_cells(row):
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = pin_to_next

    for table in tables:
        table_text = normalize(" ".join(cell.text for row in table.rows for cell in unique_cells(row)))
        if "reflection" not in table_text or "recap success criteria" not in table_text:
            continue
        for row_index, row in enumerate(table.rows):
            row_text = normalize(" ".join(cell.text for cell in unique_cells(row)))
            if allow_flow or len(row_text) > ROW_SPLIT_THRESHOLD:
                allow_row_split(row)
            else:
                set_row_cant_split(row)
            clear_row_minimum_height(row)
            if row_index < len(table.rows) - 1:
                nxt = normalize(" ".join(cell.text for cell in unique_cells(table.rows[row_index + 1])))
                for cell in unique_cells(row):
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = (
                            not allow_flow and len(nxt) <= ROW_SPLIT_THRESHOLD
                        )


# Longer than this and a row is better flowed across a page break than pushed whole.
ROW_SPLIT_THRESHOLD = 420


def allow_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for node in tr_pr.xpath("./w:cantSplit"):
        tr_pr.remove(node)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if not tr_pr.xpath("./w:cantSplit"):
        tr_pr.append(OxmlElement("w:cantSplit"))


def clear_row_minimum_height(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for height in tr_pr.xpath("./w:trHeight"):
        tr_pr.remove(height)


def split_table_before_opening(table: Table) -> None:
    split_index = None
    for index, row in enumerate(table.rows):
        row_text = normalize(" ".join(cell.text for cell in unique_cells(row)))
        if "how will you engage students in learning" in row_text:
            split_index = index
            break
    if split_index is None or split_index <= 0:
        raise TemplateError("Could not find the Opening page boundary in the template.")

    original = table._tbl
    second = deepcopy(original)
    for row in list(second.tr_lst)[:split_index]:
        second.remove(row)
    for row in list(original.tr_lst)[split_index:]:
        original.remove(row)

    separator = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_pr.append(OxmlElement("w:pageBreakBefore"))
    spacing = OxmlElement("w:spacing")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before", "0")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after", "0")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line", "20")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule", "exact")
    p_pr.append(spacing)
    separator.append(p_pr)

    parent = original.getparent()
    position = parent.index(original)
    parent.insert(position + 1, separator)
    parent.insert(position + 2, second)


def compact_trailing_paragraphs(document: Document) -> None:
    body = document._element.body
    trailing = []
    for child in reversed(list(body)):
        local_name = child.tag.rsplit("}", 1)[-1]
        if local_name == "sectPr":
            continue
        if local_name == "p" and not "".join(child.itertext()).strip():
            trailing.append(child)
            continue
        break
    for paragraph in trailing:
        body.remove(paragraph)
